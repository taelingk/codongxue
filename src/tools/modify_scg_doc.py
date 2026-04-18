from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor


TARGET = r"C:\Users\A\Desktop\心震图技术在心肺功能检测中的应用（3）.docx"
REVISED = r"D:\浏览器下载\zhuanli\SCG综述\Revised_Manuscript_SCG_Cardiopulmonary_Review (1).docx"


RED = RGBColor(192, 0, 0)


def para_text(paragraph):
    return paragraph.text.strip()


def find_paragraph(doc, startswith):
    for p in doc.paragraphs:
        if para_text(p).startswith(startswith):
            return p
    raise ValueError(f"Paragraph not found: {startswith}")


def clear_paragraph(paragraph):
    paragraph.text = ""


def set_mark_style(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = RED


def replace_paragraph(paragraph, text):
    clear_paragraph(paragraph)
    paragraph.add_run(text)
    set_mark_style(paragraph)
    return paragraph


def insert_paragraph_after(paragraph, text="", mark_red=True, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    if mark_red:
        set_mark_style(new_para)
    return new_para


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def insert_table_after(doc, paragraph, data):
    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RED
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return table


def revised_references():
    doc = Document(REVISED)
    started = False
    refs = []
    for p in doc.paragraphs:
        txt = para_text(p)
        if not txt:
            continue
        if started:
            refs.append(txt)
        elif txt == "References":
            started = True
    return refs


def main():
    doc = Document(TARGET)

    title = find_paragraph(doc, "心震图技术在心肺功能检测中的应用")
    note = insert_paragraph_after(
        title,
        "【修改标注说明】新增或替换内容已用红色字体标注，其中关键改动以“【新增】”或“【修改】”提示；参考文献已按修订稿统一更新格式。",
    )

    abstract_heading = find_paragraph(doc, "摘要：")
    abstract_body = find_paragraph(doc, "心肺功能（cardiorespiratory fitness, CRF）是评估人体有氧运动能力")
    replace_paragraph(
        abstract_body,
        "【修改摘要】Background：心肺功能（cardiorespiratory fitness, CRF）是评估心血管健康与有氧运动能力的关键生命体征，但心肺运动试验（CPET）对昂贵设备、专业场地和受试者耐受性要求较高，限制了其在社区筛查和居家监测中的推广。",
    )
    p = abstract_body
    p = insert_paragraph_after(
        p,
        "【修改摘要】Scope：本文聚焦心震图（seismocardiography, SCG）在院外心肺功能检测中的应用，围绕代谢功能（VO₂max估算）、通气功能（呼吸频率与潮气量监测）和循环功能（心输出量与每搏输出量评估）三个核心维度展开，并兼顾辅助血流动力学参数的拓展价值。",
    )
    p = insert_paragraph_after(
        p,
        "【修改摘要】Methods：参考修订稿补充了文献检索策略，系统检索PubMed、IEEE Xplore、Web of Science与Scopus数据库（2010年1月至2026年3月），纳入报告SCG估算至少一项心肺参数且提供定量性能指标的英文或中文全文研究，共纳入68篇文献。",
    )
    p = insert_paragraph_after(
        p,
        "【修改摘要】Key Findings：现有证据表明，静息SCG结合机器学习可在健康人群中实现中等精度的VO₂max估算，但模型迁移到病理人群后误差明显升高；呼吸频率提取已具有较强可行性，而潮气量估算仍更依赖多模态融合；在循环评估方面，深度学习和生理特征模型均显示出对CO/SV监测的潜力，但临床样本规模与泛化能力仍有限。",
    )
    p = insert_paragraph_after(
        p,
        "【修改摘要】Conclusions：SCG在代谢、通气和循环三个维度均已达到“方法验证”阶段，但在临床规模验证、跨人群泛化和真实场景部署方面仍处于早期探索。运动伪影、病理重塑导致的特征退化及处理流程缺乏标准化，是当前限制其临床转化的核心瓶颈。",
    )
    insert_paragraph_after(
        p,
        "【新增关键词】关键词：心震图；心肺功能；可穿戴传感；机器学习；VO₂max；心输出量；无创监测",
    )

    roadmap = find_paragraph(doc, "本综述旨在系统梳理和批判性评估SCG技术在心肺功能检测中的现有应用进展")
    replace_paragraph(
        roadmap,
        "【修改】本综述旨在系统梳理并批判性评估SCG技术在心肺功能检测中的应用进展与方法学局限。为回应审稿意见，本文在引言后补充了独立的文献检索策略章节，随后依次讨论SCG信号原理与采集算法、代谢功能、通气功能、循环功能、辅助参数评估以及方法局限与未来转化方向。",
    )

    scg_heading = find_paragraph(doc, "2. SCG技术")
    method_heading = insert_paragraph_after(
        roadmap,
        "【新增章节】文献检索策略",
        style=scg_heading.style,
    )
    p = method_heading
    p = insert_paragraph_after(
        p,
        "【新增】本文采用叙述性综述（narrative review）方法，对SCG在心肺功能评估中的代表性研究进行系统梳理，以提高综述过程的透明度并回应审稿意见中对方法学描述的要求。",
    )
    p = insert_paragraph_after(
        p,
        "【新增】数据库与时间范围：检索数据库包括PubMed/MEDLINE、IEEE Xplore、Web of Science Core Collection和Scopus，检索时间范围为2010年1月至2026年3月；对于SCG形态学与基准点定义等奠基性研究，不设起始年份限制。",
    )
    p = insert_paragraph_after(
        p,
        "【新增】检索词组合：以“seismocardiography”“seismocardiogram”“SCG”为核心词，并分别与“cardiopulmonary fitness”“VO2max”“oxygen uptake”“respiratory rate”“tidal volume”“ventilation”“cardiac output”“stroke volume”“hemodynamics”“machine learning”“deep learning”“wearable”等关键词进行布尔组合检索。",
    )
    p = insert_paragraph_after(
        p,
        "【新增】纳入与排除标准：纳入以SCG或含SCG的多模态信号为主要输入、报告至少一项代谢/通气/循环参数定量估算结果、且全文正式发表的英文或中文研究；排除仅有会议摘要而无全文、未经同行评审的预印本，以及仅以BCG为主而不含SCG数据的研究。最终纳入68篇文献用于综述分析。",
    )
    insert_paragraph_after(
        p,
        "【新增】综述类型说明：本文定位为叙述性综述而非严格意义上的系统综述，因此未采用PRISMA流程图；选择这一策略是因为现有SCG心肺功能研究在应用场景、算法类别和目标参数上高度异质，更适合进行全景式、批判性的证据整合。",
    )

    co_para = find_paragraph(doc, "在CO估算方面，Wang等人为突破CO测量对有创导管和专业影像设备的依赖")
    replace_paragraph(
        co_para,
        "【修改】在CO估算方面，Wang等人基于73名心力衰竭患者的右心导管（RHC）同步数据，构建了结合三轴SCG、ECG和BMI信息的卷积神经网络模型，用于端到端估算心输出量。该研究在低输出亚组（CO < 6 L/min）中报告平均偏差约为−0.01 L/min，提示SCG用于高风险患者连续监测具有潜力；但需要强调的是，仅报告平均偏差并不足以完全证明个体层面的一致性，文中未充分提供Bland-Altman一致性界限等信息，因此对“极高一致性”的表述应保持审慎。",
    )

    sv_para = find_paragraph(doc, "深度学习模型虽然在预测精度上表现突出，但其“黑箱”特性限制了临床可信度")
    replace_paragraph(
        sv_para,
        "【修改】深度学习模型虽然在预测精度上表现突出，但其“黑箱”特性限制了临床可信度，这也促使部分研究者转向可解释性更强的生理特征建模。Ganti等人针对先天性心脏病患者，提取SCG中的AO和AC等机械事件特征来估算每搏输出量，并以心脏磁共振（CMR）作为参考标准。该研究支持SCG用于结构异常心脏的功能评估具有可行性，但其约28%的SV误差仍提示方法距离临床独立替代尚有明显距离，更适合作为补充性监测工具，而非单独作出高风险决策的依据。",
    )

    limitation_heading = find_paragraph(doc, "SCG方法局限与挑战")
    summary_heading = insert_paragraph_after(
        find_paragraph(doc, "总体来看，这些辅助参数不作为独立终点"),
        "【新增章节】核心研究汇总表",
        style=limitation_heading.style,
    )
    p = insert_paragraph_after(summary_heading, "【新增表】表1 代谢功能相关研究汇总")
    insert_table_after(
        doc,
        p,
        [
            ["研究", "样本量", "人群", "方法", "参考标准", "关键结果", "主要局限"],
            ["Schmidt等", "数百例", "健康成人", "静息SCG特征+机器学习", "CPET", "VO₂max估算相关性较好", "仅限静息场景，泛化人群单一"],
            ["Schulenburg等", "94", "健康成人", "外部验证", "CPET", "重测信度较高", "样本规模仍有限"],
            ["Shandhi等", "约20", "健康成人", "XGBoost动态估算", "便携式气体分析", "自由活动中可追踪瞬时摄氧变化", "运动伪影明显，小样本"],
            ["Hossein等", "约50", "健康成人", "运动间歇期动能模型", "CPET", "无需极量负荷即可估算VO₂max", "依赖间歇记录，实时性受限"],
            ["Hansen等", "约60", "缺血性心脏病/心衰", "迁移修正模型", "CPET", "MAPE约29.1%", "健康模型向病理人群迁移失败"],
        ],
    )
    p = insert_paragraph_after(p, "【新增表】表2 通气功能相关研究汇总")
    insert_table_after(
        doc,
        p,
        [
            ["研究", "目标参数", "方法", "参考标准", "关键结果", "主要局限"],
            ["Naufal等", "呼吸频率", "VMD+DFA分离", "呼吸带", "提升单传感器呼吸信号分离质量", "以静息验证为主"],
            ["Chan等", "呼吸频率", "U-Net时频分离", "气流/呼吸参考", "动态条件下MAE约0.82 bpm", "需多源输入支持"],
            ["Soliman等", "潮气量", "SCG+ECG回归模型", "肺功能/气体分析", "提示SCG可为VT估算提供辅助信息", "单独SCG稳定性不足"],
            ["Imirzalioglu等", "呼吸调制特征", "TEO+机器学习", "呼吸带", "可识别呼吸相位对SCG形态的影响", "以机制探索为主"],
        ],
    )
    p = insert_paragraph_after(p, "【新增表】表3 循环功能相关研究汇总")
    insert_table_after(
        doc,
        p,
        [
            ["研究", "样本量", "目标参数", "方法", "参考标准", "关键结果", "主要局限"],
            ["Wang等", "73", "CO", "CNN端到端模型", "RHC", "低输出亚组平均偏差约−0.01 L/min", "LoA报告不足，样本仍偏小"],
            ["Ganti等", "约30", "SV", "机械特征回归", "CMR", "提示结构异常心脏中仍可估算SV", "误差约28%，临床可接受性有限"],
            ["Semiz等", "约40", "SV", "SCG-ECG融合随机森林", "超声", "多模态融合提高鲁棒性", "验证场景较局限"],
            ["Hossein等", "约50", "CO/SV", "心脏动能iK模型", "超声", "兼顾一定可解释性与相关性", "仍属间接估算指标"],
        ],
    )

    standards_para = find_paragraph(doc, "第三个核心矛盾是研究碎片化与标准化需求之间的矛盾")
    added = insert_paragraph_after(
        standards_para,
        "【新增】除算法与信号问题外，SCG走向临床应用还需面对伦理、监管与竞争技术挑战。可穿戴医疗设备若用于风险预警或治疗辅助，通常需要经历FDA或CE等监管路径验证，同时还涉及长期连续监测产生的隐私保护与数据治理问题。此外，SCG并非唯一的无创心肺监测路线，雷达心图、视频机械振动分析及其他胸壁力学传感技术也在快速发展，因此未来研究应在相同场景下开展横向比较，明确SCG真正具备优势的应用边界。",
    )

    conclusion_para = find_paragraph(doc, "综上所述，SCG正从单一机械事件检测向多维心肺功能表型拓展")
    replace_paragraph(
        conclusion_para,
        "【修改】综上所述，SCG正从单一机械事件检测走向多维心肺功能表型刻画。综合审稿意见与修订稿内容可以看出，该技术在VO₂max估算、呼吸频率监测以及CO/SV无创评估方面均已具备方法学可行性，但其临床转化仍受制于运动伪影敏感、病理状态下特征不稳定、参考标准不统一和标准化流程缺失等问题。未来更现实的推进路径，是以多模态融合和病种特异化建模为核心，在明确监管与隐私边界的前提下逐步走向真实世界应用。",
    )

    ref_heading = find_paragraph(doc, "References")
    trailing = []
    current = ref_heading._element.getnext()
    while current is not None:
        trailing.append(current)
        current = current.getnext()
    for element in trailing:
        element.getparent().remove(element)

    ref_note = insert_paragraph_after(
        ref_heading,
        "【修改说明】以下参考文献已根据修订稿统一为规范格式。",
    )
    cursor = ref_note
    for ref in revised_references():
        cursor = insert_paragraph_after(cursor, ref)

    doc.save(TARGET)


if __name__ == "__main__":
    main()
